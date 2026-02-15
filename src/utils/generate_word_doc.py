from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc():
    document = Document()
    
    # Title
    title = document.add_heading('Firebase Integration Architecture & Usage Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1
    document.add_heading('1. Executive Summary', level=1)
    p = document.add_paragraph(
        "The OCR pipeline has been upgraded to support Cloud-Native Data Synchronization. "
        "Instead of relying solely on local Excel/JSON exports, the system now automatically saves "
        "structured, patient-centric data to a Google Firebase Firestore database. "
        "This enables real-time access to medical reports from the mobile application (Flutter) "
        "and ensures data persistence and scalability."
    )
    
    # Section 2
    document.add_heading('2. System Architecture', level=1)
    p = document.add_paragraph('The system follows a "Process & Push" architecture:')
    
    items = [
        "Input: Medical Report Images/PDFs are fed into the OCR Engine.",
        "Processing: Text Extraction (Doctr & EasyOCR) and Normalization.",
        "Data Structuring: Results transformed into Patient -> Report hierarchy.",
        "Cloud Sync: Uploaded to Firebase Firestore using Admin SDK.",
        "Client Access: Flutter App subscribes to Firestore updates."
    ]
    for item in items:
        document.add_paragraph(item, style='List Bullet')

    # Section 3
    document.add_heading('3. Database Schema', level=1)
    document.add_paragraph(
        "We implemented a hierarchical schema to handle multiple reports per patient efficiently:"
    )
    
    schema = document.add_paragraph()
    schema.add_run("Collection: patients\n").bold = True
    schema.add_run("  └── Document: <Patient_ID> (10001)\n")
    schema.add_run("      ├── Name: 'BASSEL RAMY...'\n")
    schema.add_run("      └── Subcollection: reports\n").bold = True
    schema.add_run("          └── Document: <Report_ID>\n")
    schema.add_run("              ├── Date: 2026-02-14\n")
    schema.add_run("              ├── SourceFile: 'Review.pdf'\n")
    schema.add_run("              └── Results (Map):\n")
    schema.add_run("                  ├── HGB: { value: 14.5, unit: 'g/dL' }\n")
    schema.add_run("                  └── WBC: { value: 6.5, unit: '10^3/uL' }")

    # Section 4
    document.add_heading('4. Workflow Logic', level=1)
    
    document.add_heading('Step 1: Initialization', level=2)
    document.add_paragraph(
        "The FirebaseService class initializes the connection using the secure Service Account Key "
        "(serviceAccountKey.json). This grants the backend 'Admin' privileges."
    )
    
    document.add_heading('Step 2: Patient Resolution', level=2)
    document.add_paragraph(
        "For every document, the system checks if the patient exists. If yes, it updates the timestamp. If no, it creates a new record."
    )
    
    document.add_heading('Step 3: Report Upload', level=2)
    document.add_paragraph(
        "A unique Report_ID is generated. Data is uploaded to the reports subcollection, ensuring infinite scalability per patient."
    )

    # Section 5
    document.add_heading('5. Benefits', level=1)
    benefits = [
        "Patient History: Easy to query historical trends for a specific patient.",
        "Real-Time Sync: Immediate updates for the mobile app.",
        "Data Integrity: Separation of metadata and clinical values."
    ]
    for b in benefits:
        document.add_paragraph(b, style='List Bullet')

    # Save
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    output_path = os.path.join(base_dir, "docs", "Firebase_Integration_Guide.docx")
    
    document.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    create_doc()
